from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "resources" / "MMI-CORP-Question-Bank.docx"
LOGO = ROOT / "mmi-playbook-logo.png"

BLUE = RGBColor(0, 63, 120)
GREEN = RGBColor(57, 181, 74)
INK = RGBColor(23, 35, 49)
MUTED = RGBColor(91, 110, 128)
PALE_BLUE = "EAF3FA"
PALE_GREEN = "EAF7ED"
PALE_GRAY = "F7FAFC"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=INK, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullets(cell, items):
    cell.text = ""
    for idx, item in enumerate(items):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.style = "List Bullet"
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        run = p.add_run(item)
        run.font.size = Pt(8.7)
        run.font.color.rgb = INK


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.color.rgb = BLUE
    return p


def add_question_table(doc, title, purpose, groups, listen_for):
    add_heading(doc, title, 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(purpose)
    r.font.color.rgb = MUTED
    r.font.size = Pt(9.5)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.columns[0].width = Inches(1.45)
    table.columns[1].width = Inches(5.75)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Question area", True, RGBColor(255, 255, 255), 8.8)
    set_cell_text(hdr[1], "Questions to ask", True, RGBColor(255, 255, 255), 8.8)
    shade_cell(hdr[0], "003F78")
    shade_cell(hdr[1], "003F78")

    for label, questions in groups:
        row = table.add_row().cells
        shade_cell(row[0], PALE_BLUE)
        set_cell_text(row[0], label, True, BLUE, 8.7)
        add_bullets(row[1], questions)

    callout = doc.add_table(rows=1, cols=2)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.style = "Table Grid"
    callout.columns[0].width = Inches(1.45)
    callout.columns[1].width = Inches(5.75)
    row = callout.rows[0].cells
    shade_cell(row[0], PALE_GREEN)
    shade_cell(row[1], PALE_GREEN)
    set_cell_text(row[0], "Listen for", True, GREEN, 8.7)
    add_bullets(row[1], listen_for)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.2)
    for style_name in ["Heading 1", "Heading 2"]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.bold = True
        styles[style_name].font.color.rgb = BLUE

    if LOGO.exists():
        doc.add_picture(str(LOGO), width=Inches(1.7))

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("CORP Discovery Question Bank")
    r.bold = True
    r.font.size = Pt(23)
    r.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("A takeaway guide for business development calls, intro meetings, and relationship-building conversations.")
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    intro = doc.add_table(rows=1, cols=4)
    intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    intro.style = "Table Grid"
    intro_items = [
        ("C", "Contact", "Understand the person, role, influence, and communication style."),
        ("O", "Organization", "Map the business, structure, priorities, and pressure points."),
        ("R", "Recruitment", "Diagnose hiring process, challenges, decision path, and vendor usage."),
        ("P", "Projects", "Find upcoming work, skills needed, timing, and where staffing can help."),
    ]
    for idx, (letter, label, text) in enumerate(intro_items):
        cell = intro.rows[0].cells[idx]
        shade_cell(cell, PALE_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(letter)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = GREEN
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.bold = True
        r2.font.color.rgb = BLUE
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(text)
        r3.font.size = Pt(8.2)
        r3.font.color.rgb = MUTED

    tip = doc.add_paragraph()
    tip.paragraph_format.space_before = Pt(10)
    tip.paragraph_format.space_after = Pt(10)
    r = tip.add_run("How to use this: ")
    r.bold = True
    r.font.color.rgb = BLUE
    tip.add_run("Do not ask every question. Pick two or three from each area, listen carefully, then go deeper where the buyer gives you a signal.")

    sections = [
        (
            "C - Contact",
            "Learn who the person is, what they own, how they are measured, and how much influence they have over staffing decisions.",
            [
                ("Role and responsibility", [
                    "Can you walk me through your role and what you are responsible for day to day?",
                    "Which teams or functions do you support most closely?",
                    "What decisions usually come to you versus your manager or leadership team?",
                    "Where do hiring, delivery, or workforce issues show up in your role?",
                    "What does a successful quarter look like for you personally?",
                ]),
                ("Influence and decision path", [
                    "When a staffing need comes up, where do you usually get involved?",
                    "Who else weighs in before a staffing partner is approved or used?",
                    "Are you usually the person identifying the need, approving the need, or managing the work?",
                    "Who would feel the pain first if this role or project stayed uncovered?",
                    "What would make you comfortable introducing a partner to the right person internally?",
                ]),
                ("Communication style", [
                    "What is the best way to follow up with you when something is relevant?",
                    "Do you prefer quick calls, short emails, or scheduled touchpoints?",
                    "What information is useful to you, and what feels like noise?",
                    "How much detail do you like before deciding whether a conversation is worth your time?",
                    "What would make a staffing conversation feel valuable instead of transactional?",
                ]),
                ("Relationship building", [
                    "How long have you been with the organization?",
                    "What has changed most since you joined?",
                    "What are you focused on improving this year?",
                    "What types of partners have been helpful to you in the past?",
                    "What should someone understand about your team before trying to help?",
                ]),
            ],
            [
                "Role authority, personal priorities, stakeholder names, preferred communication style.",
                "Signals that they can introduce you, influence decisions, or identify pain early.",
                "Language they use to describe success, pressure, risk, or frustration.",
            ],
        ),
        (
            "O - Organization",
            "Map how the company operates, where pressure is building, and which departments are likely to need staffing support.",
            [
                ("Business and structure", [
                    "How is the organization structured across teams, locations, or business units?",
                    "Which groups are growing, changing, or under the most pressure right now?",
                    "Are decisions mostly centralized, or do departments manage hiring independently?",
                    "How does your team interact with HR, procurement, finance, or talent acquisition?",
                    "Where does outside staffing fit into the broader workforce strategy?",
                ]),
                ("Priorities and pressure", [
                    "What are the biggest business priorities for the next six to twelve months?",
                    "What initiatives are most important to leadership right now?",
                    "Where is the organization feeling capacity strain?",
                    "Are there timelines, launches, audits, go-lives, or deadlines driving urgency?",
                    "What happens operationally if roles stay open too long?",
                ]),
                ("Change signals", [
                    "Have there been recent leadership changes, reorganizations, acquisitions, or expansions?",
                    "Are any departments adding headcount or changing how work gets delivered?",
                    "Are you seeing more contract, contract-to-hire, or direct-hire needs lately?",
                    "Are budget cycles affecting hiring timing?",
                    "What has changed about hiring compared with last year?",
                ]),
                ("Vendor environment", [
                    "How does the organization currently work with staffing vendors?",
                    "Is there an MSP, VMS, preferred supplier list, or MSA process?",
                    "What does a partner need to do to earn trust here?",
                    "Where have vendors helped, and where have they fallen short?",
                    "Are there compliance, pricing, or process requirements we should understand early?",
                ]),
            ],
            [
                "Growth areas, business pressure, procurement process, account complexity, timing signals.",
                "Departments that may have repeat demand or overlooked whitespace.",
                "Barriers such as MSP/VMS rules, vendor limits, compliance, rate pressure, or budget timing.",
            ],
        ),
        (
            "R - Recruitment",
            "Diagnose the hiring motion: what they need, why it is hard, who is involved, and where the current process breaks down.",
            [
                ("Current hiring landscape", [
                    "What does your current hiring landscape look like?",
                    "Which roles have been hardest to fill recently?",
                    "Where are you seeing the biggest skill gaps?",
                    "Are openings mostly growth-related, replacement, project-based, or backfill?",
                    "What roles are most important to fill quickly?",
                ]),
                ("Process and decision path", [
                    "What does the hiring process look like from opening a role to making an offer?",
                    "Who owns requirements, interview feedback, approvals, and final decision?",
                    "How many interview steps are typical?",
                    "Where does the process usually slow down?",
                    "How quickly can the team move when the right candidate is identified?",
                ]),
                ("Quality and fit", [
                    "What separates a strong candidate from an average one for your team?",
                    "What skills are non-negotiable versus nice to have?",
                    "What soft skills or work style traits matter most?",
                    "What causes candidates to fail in the interview process?",
                    "What does someone need to understand about the environment to succeed after day one?",
                ]),
                ("Market and delivery friction", [
                    "Are you seeing enough qualified candidates from current channels?",
                    "Are compensation, location, remote policy, or speed affecting candidate interest?",
                    "How often do candidates drop out or accept competing offers?",
                    "What feedback do candidates give about the opportunity?",
                    "Where could a staffing partner make the biggest difference?",
                ]),
                ("Partner expectations", [
                    "What does a good staffing partner do differently for your team?",
                    "What information do you need upfront to trust a submission?",
                    "How do you prefer to receive candidate updates?",
                    "What would make you want to use a partner again?",
                    "If we found the right profile, what would the next step be?",
                ]),
            ],
            [
                "Hard-to-fill roles, real urgency, decision makers, process gaps, candidate fit criteria.",
                "Pain around speed, quality, submittal relevance, compensation, interview discipline, or competing offers.",
                "Clear next step: intro meeting, job order, stakeholder introduction, or follow-up with market data.",
            ],
        ),
        (
            "P - Projects",
            "Find upcoming work that creates staffing demand, then connect talent strategy to timing, skills, and business impact.",
            [
                ("Upcoming initiatives", [
                    "What projects or initiatives are coming up that may require additional talent?",
                    "Are there implementations, upgrades, migrations, audits, launches, or transformations planned?",
                    "Which projects are highest priority for leadership?",
                    "What work is currently at risk because of capacity or skill gaps?",
                    "Where do you expect workload to increase over the next quarter?",
                ]),
                ("Skills and team needs", [
                    "What skills will be required to deliver the project successfully?",
                    "Are those skills already on the team, or will you need outside support?",
                    "Do you need individual contributors, team leads, project managers, or niche specialists?",
                    "What experience would make someone productive quickly?",
                    "Are certifications, domain experience, systems knowledge, or compliance requirements important?",
                ]),
                ("Timing and urgency", [
                    "What is the target start date or deadline?",
                    "What milestones could be affected if the right people are not in place?",
                    "When would you need to see candidates to stay on track?",
                    "How long can the team operate without this support?",
                    "What would trigger a decision to bring in contract help?",
                ]),
                ("Budget and approval", [
                    "Is budget already allocated, or would approval be needed?",
                    "Who approves outside project support?",
                    "Is the work better suited for contract, contract-to-hire, direct hire, or a project team?",
                    "Are there rate, tenure, location, or compliance constraints?",
                    "What information would help leadership approve support faster?",
                ]),
                ("Expansion opportunity", [
                    "If we helped on this project, where else could similar support be useful?",
                    "Are other teams facing the same skill or capacity challenge?",
                    "Who else is responsible for related initiatives?",
                    "Would it be useful to compare the market for these skills before the project ramps up?",
                    "What would a successful partnership look like by the end of the project?",
                ]),
            ],
            [
                "Project timing, business consequence, skills needed, budget state, approval path.",
                "Signals for contract demand, specialist needs, team augmentation, or repeat project work.",
                "Opportunities to introduce market intelligence, talent availability, or a scoped follow-up meeting.",
            ],
        ),
    ]

    for idx, args in enumerate(sections):
        if idx == 2:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_question_table(doc, *args)

    doc.add_page_break()
    add_heading(doc, "Quick call flow", 1)
    flow = doc.add_table(rows=1, cols=4)
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    flow.style = "Table Grid"
    flow_items = [
        ("Open", "State who you are, why you are calling, and ask permission for one question."),
        ("Diagnose", "Use CORP to understand person, company, hiring motion, and project pressure."),
        ("Connect", "Tie MMI value to the specific pain or opportunity they shared."),
        ("Advance", "Secure a next step with owner, date, purpose, and useful follow-up."),
    ]
    for idx, (label, text) in enumerate(flow_items):
        cell = flow.rows[0].cells[idx]
        shade_cell(cell, PALE_BLUE if idx % 2 == 0 else PALE_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = BLUE
        r.font.size = Pt(11)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(text)
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = INK

    add_heading(doc, "Notes from the conversation", 1)
    notes = doc.add_table(rows=6, cols=2)
    notes.alignment = WD_TABLE_ALIGNMENT.CENTER
    notes.style = "Table Grid"
    labels = ["Contact", "Organization", "Recruitment", "Projects", "Pain / trigger", "Next step"]
    for idx, label in enumerate(labels):
        left, right = notes.rows[idx].cells
        shade_cell(left, PALE_BLUE)
        set_cell_text(left, label, True, BLUE, 8.8)
        set_cell_text(right, "", False, INK, 8.8)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)
